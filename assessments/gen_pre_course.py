"""Generate assessments/pre_course.docx and assessments/pre_course.xlsx."""

import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

# ── QUESTIONS ─────────────────────────────────────────────────────────────────
# 14 conceptual questions, no coding required, entry level
# Topics span the full course arc at pre-knowledge level.
# ESLATMA: keyinchalik to'liq test bankiga birlashtiriladi.
questions = [
    {
        "no": 1, "topic": "NLP asoslari",
        "q": "Quyidagilardan qaysi biri NLP vazifasi EMAS?",
        "opts": [
            "A) Matnni sentiment bo'yicha tasniflash",
            "B) Rasmdan ob'ektlarni aniqlash",
            "C) Matndan shaxs nomlarini ajratib olish",
            "D) So'zlarni morfologik tahlil qilish",
        ],
        "ans": "B",
        "expl": "Rasm ob'ektlarni aniqlash Computer Vision vazifasi.",
    },
    {
        "no": 2, "topic": "Tokenizatsiya",
        "q": (
            "\"O'zbekistonda\" so'zi so'z-tokenizatsiyadan o'tkazilganda "
            "qaysi natija to'g'ri?"
        ),
        "opts": [
            "A) [\"O'zbekistonda\"]  — 1 token",
            "B) [\"O'zbekiston\", \"da\"] — 2 token",
            "C) [\"O'z\", \"bek\", \"iston\", \"da\"] — 4 token",
            "D) [\"o\", \"'\", \"zbekistonda\"] — 3 token",
        ],
        "ans": "A",
        "expl": "So'z tokenizatsiyasi butun so'zni bir token qiladi.",
    },
    {
        "no": 3, "topic": "Bag-of-Words",
        "q": "Bag-of-Words modeli qaysi ma'lumotni YO'QOTADI?",
        "opts": [
            "A) So'zlarning chastotasini",
            "B) So'zlarning tartibini (pozitsiyasini)",
            "C) So'zlarning hujjatda mavjudligini",
            "D) Hujjat uzunligini",
        ],
        "ans": "B",
        "expl": "BoW so'z tartibi haqida ma'lumot saqlamaydi.",
    },
    {
        "no": 4, "topic": "TF-IDF",
        "q": "Qaysi so'zning TF-IDF qiymati eng YUQORI bo'ladi?",
        "opts": [
            "A) Hamma hujjatda ko'p uchraydigan so'z",
            "B) Faqat bitta hujjatda ko'p uchraydigan noyob so'z",
            "C) Hech bir hujjatda uchramaydigan so'z",
            "D) Barcha hujjatlarda bir marta uchraydigan so'z",
        ],
        "ans": "B",
        "expl": "TF-IDF = TF x IDF. Noyob so'z IDF ni oshiradi.",
    },
    {
        "no": 5, "topic": "Metrikalar",
        "q": (
            "100 ta musbat namunadan model 80 tasini to'g'ri topdi, "
            "20 tasini o'tkazib yubordi. Recall qancha?"
        ),
        "opts": ["A) 0.80", "B) 0.20", "C) 1.00", "D) Hisoblab bo'lmaydi"],
        "ans": "A",
        "expl": "Recall = 80 / (80+20) = 0.80.",
    },
    {
        "no": 6, "topic": "Embeddinglar",
        "q": "Word2Vec ning BoW dan asosiy afzalligi nima?",
        "opts": [
            "A) Har doim tezroq hisoblaydi",
            "B) Semantik o'xshashlikni aks ettiradi",
            "C) So'z tartibini saqlaydi",
            "D) Grammatik qoidalarni kodlaydi",
        ],
        "ans": "B",
        "expl": "Word2Vec semantik yaqin so'zlarga yaqin vektorlar beradi.",
    },
    {
        "no": 7, "topic": "Neyron tarmoqlar",
        "q": "Neyron tarmoq nima uchun o'quv ma'lumotiga muhtoj?",
        "opts": [
            "A) Parametrlarini (og'irliklarini) o'rnatish uchun",
            "B) Faqat test natijalarini ko'rsatish uchun",
            "C) Foydalanuvchi interfeysini yaratish uchun",
            "D) Lug'at hajmini aniqlash uchun",
        ],
        "ans": "A",
        "expl": "Gradient orqali parametrlar o'quv ma'lumotidan o'rganiladi.",
    },
    {
        "no": 8, "topic": "RNN",
        "q": "RNN oddiy to'liq ulangan tarmoqdan qaysi xossasi bilan farq qiladi?",
        "opts": [
            "A) Oldingi vaqt qadamidagi yashirin holatni keyingisiga uzatadi",
            "B) Har doim ko'proq qatlamga ega",
            "C) Faqat tasvir bilan ishlaydi",
            "D) Gradient hisoblash shart emas",
        ],
        "ans": "A",
        "expl": "h_t = f(h_{t-1}, x_t) — ketma-ketlik bo'yicha yashirin holat.",
    },
    {
        "no": 9, "topic": "Transfer learning",
        "q": "\"Fine-tuning\" nima degani?",
        "opts": [
            "A) Modelni noldan o'qitish",
            "B) Oldindan o'qitilgan modelni yangi vazifaga moslash",
            "C) Modelni siqib fayl hajmini kamaytirish",
            "D) Ma'lumotlarni tozalash jarayoni",
        ],
        "ans": "B",
        "expl": "BERT kabi model yangi task uchun oz ma'lumotda moslashtiriladi.",
    },
    {
        "no": 10, "topic": "O'zbek NLP",
        "q": "O'zbek tili NLP uchun qaysi xususiyat asosiy qiyinchilik tug'diradi?",
        "opts": [
            "A) Kirilincha va Lotin yozuvi",
            "B) Agglutinativ morfologiya va ma'lumotlar kamligi",
            "C) Faqat SOV gap tartibi",
            "D) Inglizcha so'zlar ko'pligi",
        ],
        "ans": "B",
        "expl": "Agglutinatsiya vocabulary ni kattalashtiradi; labeled data kam.",
    },
    {
        "no": 11, "topic": "API",
        "q": "REST API nima qiladi?",
        "opts": [
            "A) Faqat tasvirlarni qayta ishlaydi",
            "B) HTTP orqali dasturlar o'rtasida ma'lumot almashishga imkon beradi",
            "C) Ma'lumotlarni faqat o'qiydi, o'zgartira olmaydi",
            "D) Faqat mahalliy tarmoqda ishlaydi",
        ],
        "ans": "B",
        "expl": "REST API — HTTP GET/POST/PUT orqali ma'lumot almashish standarti.",
    },
    {
        "no": 12, "topic": "Ma'lumotlar",
        "q": "Mashina o'rganishida train/test split nima uchun kerak?",
        "opts": [
            "A) Ma'lumotlar hajmini kamaytirish uchun",
            "B) Modelni ko'rmagan ma'lumotlarda baholash uchun",
            "C) Ma'lumotlarni shifrlash uchun",
            "D) Model tezligini oshirish uchun",
        ],
        "ans": "B",
        "expl": "Test haqiqiy ishlash qobiliyatini baholaydi (overfitting aniqlash).",
    },
    {
        "no": 13, "topic": "RAG",
        "q": "LLM \"gallyutsinatsiya\" qilishi degani nima?",
        "opts": [
            "A) Model juda sekin ishlashi",
            "B) Model noto'g'ri yoki uydirma ma'lumotni qaytarishi",
            "C) Model manba tilini noto'g'ri aniqlashi",
            "D) Model grammatik xato qilishi",
        ],
        "ans": "B",
        "expl": "LLM bilmagan narsani bilaman deb ishonchli yolg'on berishi.",
    },
    {
        "no": 14, "topic": "Python",
        "q": (
            "Python da quyidagi kod nima chiqaradi?\n"
            "words = [\"nlp\", \"juda\", \"qiziq\"]\n"
            "print(len([w for w in words if len(w) > 3]))"
        ),
        "opts": ["A) 0", "B) 1", "C) 2", "D) 3"],
        "ans": "B",
        "expl": 'Faqat "qiziq" (5 belgi) > 3 shartni o\'tadi. Natija: 1.',
    },
]

# Fix the dict syntax issue for q13
questions[12]["opts"] = [
    "A) Model juda sekin ishlashi",
    "B) Model noto'g'ri yoki uydirma ma'lumotni ishonch bilan qaytarishi",
    "C) Model manba tilini noto'g'ri aniqlashi",
    "D) Model grammatik xato qilishi",
]

# ── DOCX ──────────────────────────────────────────────────────────────────────
doc = Document()

# Margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.2)
    section.right_margin = Inches(1.2)

doc.styles["Normal"].font.name = "Times New Roman"
doc.styles["Normal"].font.size = Pt(12)

h1 = doc.add_heading("Kirish Diagnostik Testi", level=1)
h1.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Kurs: Tabiiy tilni qayta ishlash (NLP) | VMQ 425-son")
r.bold = True

doc.add_paragraph("Sana: 15-iyun 2026  |  Vaqt: 30 daqiqa  |  Kodlash talab qilinmaydi")

note = doc.add_paragraph()
note.add_run(
    "ESLATMA: Ushbu test keyinchalik to'liq test bankiga birlashtiriladi "
    "(kurs oxiridagi yakuniy test bilan o'sishni o'lchash uchun). "
    "Natijalar o'qitish maqsadida — baholash uchun emas."
)
note.runs[0].italic = True

doc.add_paragraph()
h2 = doc.add_heading("Ko'rsatma", level=2)
doc.add_paragraph(
    "Har savol uchun TO'G'RI javob harfini (A, B, C yoki D) belgilang. "
    "Hamma savollarga javob bering — bitta to'g'ri javob bor."
)
doc.add_paragraph()

for q in questions:
    p_q = doc.add_paragraph()
    r_num = p_q.add_run(f"{q['no']}. ")
    r_num.bold = True
    r_topic = p_q.add_run(f"[{q['topic']}]  ")
    r_topic.bold = True
    r_topic.italic = True
    p_q.add_run(q["q"])

    for opt in q["opts"]:
        p_o = doc.add_paragraph(opt)
        p_o.paragraph_format.left_indent = Inches(0.4)

    doc.add_paragraph()

out_docx = os.path.join(os.path.dirname(__file__), "pre_course.docx")
doc.save(out_docx)
print(f"OK: {out_docx}")

# ── XLSX ──────────────────────────────────────────────────────────────────────
wb = openpyxl.Workbook()
ws = wb.active
ws.title = "Answer Key"

thin = Side(style="thin", color="AAAAAA")
bdr  = Border(left=thin, right=thin, top=thin, bottom=thin)
hfill = PatternFill(fill_type="solid", fgColor="0E2399")
hfont = Font(color="FFFFFF", bold=True, name="Calibri", size=11)
efill = PatternFill(fill_type="solid", fgColor="EEF0FB")

headers = ["#", "Mavzu", "Savol (qisqacha)", "Javob", "Tushuntirish"]
widths  = [4,   20,      60,                 8,       65]
for col, (h, w) in enumerate(zip(headers, widths), 1):
    c = ws.cell(row=1, column=col, value=h)
    c.font = hfill and hfont
    c.fill = hfill
    c.border = bdr
    c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
    ws.column_dimensions[c.column_letter].width = w
ws.row_dimensions[1].height = 28

for i, q in enumerate(questions):
    row = i + 2
    short_q = q["q"][:85] + ("..." if len(q["q"]) > 85 else "")
    vals = [q["no"], q["topic"], short_q, q["ans"], q["expl"]]
    fill = efill if i % 2 == 0 else PatternFill(fill_type="solid", fgColor="FFFFFF")
    for col, val in enumerate(vals, 1):
        c = ws.cell(row=row, column=col, value=val)
        c.fill = fill
        c.border = bdr
        c.alignment = Alignment(vertical="top", wrap_text=True)
        if col == 4:
            c.font = Font(bold=True, size=12)
            c.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[row].height = 38

ws2 = wb.create_sheet("Meta")
meta_rows = [
    ("Kurs",           "Tabiiy tilni qayta ishlash (NLP) | VMQ 425-son"),
    ("Test turi",      "Kirish diagnostik (pre-course)"),
    ("Sana",           "15-iyun 2026"),
    ("Savollar soni",  len(questions)),
    ("Eslatma",        "Keyinchalik to'liq test bankiga birlashtiriladi."),
    ("Varaq hujjati",  "pre_course.docx"),
]
for r, (k, v) in enumerate(meta_rows, 1):
    ws2.cell(row=r, column=1, value=k).font = Font(bold=True)
    ws2.cell(row=r, column=2, value=v)
ws2.column_dimensions["A"].width = 22
ws2.column_dimensions["B"].width = 60

out_xlsx = os.path.join(os.path.dirname(__file__), "pre_course.xlsx")
wb.save(out_xlsx)
print(f"OK: {out_xlsx}")
